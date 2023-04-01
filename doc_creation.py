from docx import Document
from datetime import datetime

# I can create the schema of the 'CUSTOMERS' variable below
# using another script I have to read csv's

CUSTOMERS = [
    {
        "name": 'Bill Smith',
        'charges': [
            {'cost': '100', "date": "3-1-2023", 'procedure': 'bill_action_1'},
            {'cost': '200', "date": "3-2-2023", 'procedure': 'bill_action_2'},
            {'cost': '300', "date": "3-3-2023", 'procedure': 'bill_action_3'}
        ]
    },
    {
        "name": 'Sally Smith',
        'charges': [
            {'cost': '150', "date": "4-1-2023", 'procedure': 'sally_action_1'},
            {'cost': '250', "date": "4-2-2023", 'procedure': 'sally_action_2'},
            {'cost': '350', "date": "4-3-2023", 'procedure': 'sally_action_3'},
            {'cost': '450', "date": "4-4-2023", 'procedure': 'sally_action_4'},
            {'cost': '550', "date": "4-5-2023", 'procedure': 'sally_action_5'},
            {'cost': '650', "date": "4-6-2023", 'procedure': 'sally_action_6'},
            {'cost': '750', "date": "4-7-2023", 'procedure': 'sally_action_7'}
        ]
    }
]

# Set this value to adjust how many
# bill line items are printed per page
RECORDS_PER_PAGE = 3



def main():
    '''This is the main function of the program'''
    document = Document()
    
    for customer in CUSTOMERS:
        for charge_group in get_charge_group_lists(customer, RECORDS_PER_PAGE):
            new_page(customer, charge_group, document)

    current_time = datetime.now().strftime("%d%h_%H_%M_%S")
    document.save(f'{current_time}_Billing_Data.docx')




def new_page(cust_data, charge_group, document):
    ''' This function writes text to the new page'''

    title = document.add_paragraph('Dear Insurance company,')

    name = cust_data['name']
    p = document.add_paragraph(f'The customer {name} has the following claim charges:')

    for charge in charge_group:
        cost = charge['cost']
        date = charge['date']
        procedure = charge['procedure']
        document.add_paragraph(f"{date}\nProcedure: {procedure}\nCost: {cost}")

    document.add_page_break()

def get_charge_group_lists(cust_data, records_per_page):
    temp = []
    output = []
    for count, charge in enumerate(cust_data['charges']):
        if count == 0:
            temp.append(charge)
        
        elif count % records_per_page != 0:
            temp.append(charge)

        else:
            output.append(temp)
            temp = [charge]

        if len(cust_data['charges']) == count+1:
            output.append(temp)

    return output


if __name__ == "__main__":
    main()
